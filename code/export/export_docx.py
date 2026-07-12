from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Inches, Pt

from common import (
    IMAGE_RE,
    MARKDOWN_LINK_RE,
    WINDOWS_CONSOLAS,
    build_config,
    choose_writable_output_path,
    is_formula_like_block,
    is_probably_local_target,
    normalize_links,
    parse_args_with_version,
    read_text,
    resolve_repo_target,
)
from docx_math_upgrade import (
    append_formula_omml,
    append_matrix_formula,
    build_formula_omml,
    is_matrix_formula,
    latex_to_word_linear,
    normalize_formula_block,
    register_linear_formula,
    upgrade_docx_linear_formulas,
)
from render_assets import render_mermaid_to_png, stable_render_path


TABLE_DIVIDER_RE = re.compile(r"^\|\s*---")
LOCAL_TABLE_LINK_RE = re.compile(r"^见[:：]?\s*\[(?P<label>[^\]]+)\]\((?P<target>[^)]+)\)[。.]?$")
SOURCE_LINE_RE = re.compile(r"^源文件：`(?P<target>[^`]+)`")
FRONT_MATTER_H2 = {"版本说明", "结构说明", "章节索引", "配套文件", "入口文件职责补充说明"}


INLINE_TOKEN_RE = re.compile(r"\\\((?P<math>.+?)\\\)|`(?P<code>[^`]+)`")


def formula_font_name() -> str:
    if WINDOWS_CONSOLAS.exists():
        return "Consolas"
    return "仿宋"


def apply_run_font(run, font_name: str, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:cs"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def apply_font_to_paragraph(paragraph, font_name: str, size: float, bold: bool | None = None) -> None:
    for run in paragraph.runs:
        apply_run_font(run, font_name, size=size, bold=bold)


def set_cell_text(cell, text: str, font_name: str) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            apply_run_font(run, font_name, size=10.5)


def add_code_paragraph(document: Document, block: list[str], font_name: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.left_indent = Inches(0.2)
    run = paragraph.add_run("\n".join(block))
    apply_run_font(run, font_name, size=10)


def add_formula_paragraph(document: Document, block: list[str], font_name: str, formulas) -> None:
    formula_text = normalize_formula_block(block)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.widow_control = True
    try:
        append_formula_omml(paragraph, formula_text)
        return
    except Exception:
        if is_matrix_formula(formula_text):
            raise
    run = paragraph.add_run(register_linear_formula(formulas, formula_text))
    apply_run_font(run, font_name, size=12)


def sanitize_inline_images(text: str) -> str:
    def repl(match: re.Match[str]) -> str:
        alt = match.group(1).strip()
        target = match.group(2).strip()
        label = alt or Path(target).name
        return f"[鍥剧墖: {label}]"

    return IMAGE_RE.sub(repl, text)


def looks_like_math_code(text: str) -> bool:
    value = text.strip()
    if not value:
        return False
    if any(token in value for token in ("://", ".md", ".mmd", ".csv", ".py", "docs/", "assets/", "research/", "data/", "output/")):
        return False
    if re.search(r"\\[A-Za-z]+", value):
        return True
    if re.fullmatch(r"[A-Za-zΑ-Ωα-ω]", value):
        return True
    if re.fullmatch(r"[A-Za-zΑ-Ωα-ω](?:_[A-Za-z0-9{}+\-]+)?(?:\^[A-Za-z0-9{}*+\-]+)?", value):
        return True
    if any(char in value for char in ("_", "^", "{", "}")):
        return True
    if re.search(r"\b(?:SE|SO)\(\d+\)", value):
        return True
    if re.fullmatch(r"[A-Za-zΑ-Ωα-ω][A-Za-z0-9]{0,2}\([^()]+\)", value):
        return True
    if re.search(r"[A-Za-zΑ-Ωα-ω]\([^)]+\)", value) and any(op in value for op in ("=", "+", "-", "/", "|")):
        return True
    return False


def iter_inline_fragments(text: str):
    cursor = 0
    for match in INLINE_TOKEN_RE.finditer(text):
        if match.start() > cursor:
            yield ("text", text[cursor:match.start()])
        if match.group("math") is not None:
            yield ("math", match.group("math").strip())
        else:
            content = match.group("code")
            yield ("math" if looks_like_math_code(content) else "code", content)
        cursor = match.end()
    if cursor < len(text):
        yield ("text", text[cursor:])


def iter_markdown_fragments(text: str):
    sanitized = sanitize_inline_images(text)
    cursor = 0
    for match in MARKDOWN_LINK_RE.finditer(sanitized):
        if match.start() > cursor:
            yield from iter_inline_fragments(sanitized[cursor:match.start()])
        label = match.group(1).strip()
        target = match.group(2).strip()
        if is_probably_local_target(target):
            yield from iter_inline_fragments(label)
        else:
            yield from iter_inline_fragments(label)
        cursor = match.end()
    if cursor < len(sanitized):
        yield from iter_inline_fragments(sanitized[cursor:])


def append_markdown_inlines(paragraph, text: str, font_name: str, mono_font: str, size: float) -> None:
    for kind, content in iter_markdown_fragments(text):
        if not content:
            continue
        if kind == "text":
            run = paragraph.add_run(content)
            apply_run_font(run, font_name, size=size)
            continue
        if kind == "code":
            run = paragraph.add_run(content)
            apply_run_font(run, mono_font, size=size)
            continue
        try:
            paragraph._element.append(build_formula_omml(content))
        except Exception:
            run = paragraph.add_run(latex_to_word_linear(content))
            apply_run_font(run, mono_font, size=size)


def add_front_matter_heading(document: Document, text: str, font_name: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    apply_run_font(run, font_name, size=15, bold=True)


def add_table(document: Document, rows: list[list[str]], font_name: str) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.autofit = True
    for row_idx, row in enumerate(rows):
        for col_idx in range(cols):
            value = row[col_idx] if col_idx < len(row) else ""
            set_cell_text(table.cell(row_idx, col_idx), value, font_name)
            if row_idx == 0:
                for p in table.cell(row_idx, col_idx).paragraphs:
                    for run in p.runs:
                        run.bold = True
    document.add_paragraph("")


def add_manual_numbered_item(document: Document, index: int, text: str, font_name: str, mono_font: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.35)
    paragraph.paragraph_format.first_line_indent = Inches(-0.22)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.3
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    append_markdown_inlines(paragraph, f"{index}. {text}", font_name, mono_font, 11.5)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines):
        line = lines[i]
        if not line.strip().startswith("|"):
            break
        if TABLE_DIVIDER_RE.match(line.strip()):
            i += 1
            continue
        cells = [normalize_links(cell.strip()) for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def extract_table_from_markdown(path: Path) -> list[list[str]]:
    rows: list[list[str]] = []
    lines = read_text(path).splitlines()
    for idx, line in enumerate(lines):
        if line.strip().startswith("|"):
            rows, _ = parse_table(lines, idx)
            if rows:
                return rows
    return rows


def apply_document_style(document: Document, version: str, source: str, font_name: str) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal = document.styles["Normal"]
    normal.font.name = font_name
    normal._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    normal._element.rPr.rFonts.set(qn("w:cs"), font_name)
    normal.font.size = Pt(11.5)
    normal.font.color.rgb = RGBColor(0, 0, 0)

    for style_name, size in [("Heading 1", 18), ("Heading 2", 15), ("Heading 3", 13), ("Heading 4", 11.5)]:
        style = document.styles[style_name]
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn("w:ascii"), font_name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style._element.rPr.rFonts.set(qn("w:cs"), font_name)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)

    header = section.header.paragraphs[0]
    header.text = f"具身智能行业和技术发展调研报告 {version}"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        apply_run_font(run, font_name, size=9.5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number_field(footer)


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_image(document: Document, image_path: Path, width_inches: float = 6.0) -> None:
    document.add_picture(str(image_path), width=Inches(width_inches))
    last = document.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("")


def render_markdown_to_docx(markdown_text: str, output_path: Path, config) -> None:
    document = Document()
    font_name = "仿宋"
    mono_font = formula_font_name()
    formulas = []
    apply_document_style(document, config.version, config.source, font_name)

    lines = markdown_text.splitlines()
    i = 0
    in_code = False
    code_kind = ""
    code_lines: list[str] = []
    in_front_matter = True

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        if stripped.startswith("```"):
            if in_code:
                if code_kind == "math" or is_formula_like_block("\n".join(code_lines)):
                    add_formula_paragraph(document, code_lines, mono_font, formulas)
                elif code_kind == "mermaid":
                    pass
                else:
                    add_code_paragraph(document, code_lines, mono_font)
                code_lines = []
                in_code = False
                code_kind = ""
            else:
                in_code = True
                code_kind = stripped.strip("`").strip()
            i += 1
            continue

        if in_code:
            code_lines.append(raw.rstrip())
            i += 1
            continue

        if not stripped or stripped.startswith("<!--"):
            i += 1
            continue

        if stripped == "---":
            document.add_page_break()
            i += 1
            continue

        if stripped == r"\[":
            math_lines: list[str] = []
            i += 1
            while i < len(lines) and lines[i].strip() != r"\]":
                math_lines.append(lines[i].rstrip())
                i += 1
            if i < len(lines) and lines[i].strip() == r"\]":
                i += 1
            add_formula_paragraph(document, math_lines, mono_font, formulas)
            continue

        if stripped.startswith("|"):
            rows, next_i = parse_table(lines, i)
            add_table(document, rows, font_name)
            i = next_i
            continue

        if stripped.startswith("# "):
            text = normalize_links(stripped[2:].strip())
            if "调研报告-v" in text:
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                apply_run_font(run, font_name, size=20, bold=True)
            else:
                in_front_matter = False
                p = document.add_paragraph(text, style="Heading 1")
                apply_font_to_paragraph(p, font_name, 18, bold=True)
            i += 1
            continue

        if stripped.startswith("## "):
            heading_text = normalize_links(stripped[3:].strip())
            if in_front_matter and heading_text in FRONT_MATTER_H2:
                add_front_matter_heading(document, heading_text, font_name)
            else:
                p = document.add_paragraph(heading_text, style="Heading 2")
                apply_font_to_paragraph(p, font_name, 15, bold=True)
            i += 1
            continue

        if stripped.startswith("### "):
            p = document.add_paragraph(normalize_links(stripped[4:].strip()), style="Heading 3")
            apply_font_to_paragraph(p, font_name, 13, bold=True)
            i += 1
            continue

        if stripped.startswith("#### "):
            p = document.add_paragraph(normalize_links(stripped[5:].strip()), style="Heading 4")
            apply_font_to_paragraph(p, font_name, 11.5, bold=True)
            i += 1
            continue

        source_match = SOURCE_LINE_RE.match(stripped)
        if source_match:
            target = source_match.group("target")
            source_path = resolve_repo_target(target)
            if source_path and source_path.suffix.lower() == ".mmd" and source_path.exists():
                rendered = stable_render_path(config.rendered_assets_dir, "diagrams", str(source_path), "png")
                render_mermaid_to_png(source_path, rendered)
                add_image(document, rendered)
            i += 1
            continue

        table_link_match = LOCAL_TABLE_LINK_RE.match(stripped)
        if table_link_match:
            target = table_link_match.group("target")
            table_path = resolve_repo_target(target)
            if table_path and table_path.exists():
                rows = extract_table_from_markdown(table_path)
                add_table(document, rows, font_name)
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            list_items: list[str] = []
            while i < len(lines):
                numbered = lines[i].strip()
                if not re.match(r"^\d+\.\s+", numbered):
                    break
                list_items.append(re.sub(r"^\d+\.\s+", "", numbered))
                i += 1
            for item_index, item_text in enumerate(list_items, 1):
                add_manual_numbered_item(document, item_index, item_text, font_name, mono_font)
            continue

        if stripped.startswith("- "):
            p = document.add_paragraph(style="List Bullet")
            append_markdown_inlines(p, stripped[2:].strip(), font_name, mono_font, 11.5)
            i += 1
            continue

        if stripped.startswith("> "):
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            append_markdown_inlines(p, stripped[2:].strip(), font_name, mono_font, 11.5)
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            peek = lines[i].strip()
            if not peek:
                break
            if (
                peek.startswith("#")
                or peek.startswith("```")
                or peek == r"\["
                or peek.startswith("|")
                or peek == "---"
                or peek.startswith("- ")
                or re.match(r"^\d+\.\s+", peek)
                or peek.startswith("> ")
                or SOURCE_LINE_RE.match(peek)
                or LOCAL_TABLE_LINK_RE.match(peek)
            ):
                break
            paragraph_lines.append(peek)
            i += 1

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.35
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        append_markdown_inlines(paragraph, " ".join(paragraph_lines), font_name, mono_font, 11.5)

    document.save(output_path)
    upgrade_docx_linear_formulas(output_path, formulas)


def main() -> None:
    args = parse_args_with_version()
    config = build_config(args.version, args.source)
    markdown_text = read_text(config.bundle_path)
    output_path = choose_writable_output_path(config.docx_path)
    render_markdown_to_docx(markdown_text, output_path, config)
    print(f"DOCX written: {output_path}")


if __name__ == "__main__":
    main()
